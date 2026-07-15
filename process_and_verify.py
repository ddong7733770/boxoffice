import os
import re
from datetime import datetime
import docx
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Blue color for link
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0073E6')
    rPr.append(color)
    
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    # Font name setting
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:hint'), 'eastAsia')
    rFonts.set(qn('w:ascii'), 'Malgun Gothic')
    rFonts.set(qn('w:hAnsi'), 'Malgun Gothic')
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    rPr.append(rFonts)
    
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="DCDCDC"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="DCDCDC"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="DCDCDC"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="DCDCDC"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="DCDCDC"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="DCDCDC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def process_text_runs(paragraph, text):
    # Parse markdown bold (**) and links ([text](url))
    # Returns runs for paragraph
    pattern = r'(\*\*.*?\*\*|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            # Bold
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.font.name = 'Malgun Gothic'
            run.font.bold = True
        elif part.startswith('[') and '](' in part and part.endswith(')'):
            # Hyperlink
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                link_text = link_match.group(1)
                link_url = link_match.group(2)
                add_hyperlink(paragraph, link_text, link_url)
        else:
            # Plain text
            run = paragraph.add_run(part)
            run.font.name = 'Malgun Gothic'

def add_docx_paragraph(doc, text, style_type="Normal"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    
    if style_type == "H1":
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(18)
        run.font.bold = True
    elif style_type == "H2":
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(14)
        run.font.bold = True
    elif style_type == "H3":
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(11)
        run.font.bold = True
    else:
        # Normal paragraph parsing links and bold
        process_text_runs(p, text)
    return p

def build_docx_table(doc, t_lines):
    rows_data = []
    max_cols = 0
    
    for line in t_lines:
        if re.match(r'^\s*\|?\s*:\s*[-=]+', line.strip()) or '---' in line:
            continue
        clean_line = line.strip()
        if clean_line.startswith("|"):
            clean_line = clean_line[1:]
        if clean_line.endswith("|"):
            clean_line = clean_line[:-1]
        
        cols = [c.strip() for c in clean_line.split("|")]
        if len(cols) > max_cols:
            max_cols = len(cols)
        rows_data.append(cols)
        
    if not rows_data or max_cols == 0:
        return
        
    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    set_table_borders(table)
    
    for r_idx, row in enumerate(rows_data):
        # Stylize header row background
        if r_idx == 0:
            for cell in table.rows[0].cells:
                set_cell_background(cell, "F2F2F2")
                
        for c_idx, val in enumerate(row):
            if c_idx >= max_cols:
                break
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            # Make header text bold
            if r_idx == 0:
                run = p.add_run()
                run.font.bold = True
                run.font.name = 'Malgun Gothic'
                # Render content
                process_text_runs(p, val)
            else:
                process_text_runs(p, val)

def convert_md_to_docx(md_content, output_docx_path):
    doc = docx.Document()
    
    # Set default document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(40)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(40)
        section.right_margin = Pt(40)
        
    lines = md_content.split('\n')
    in_table = False
    table_lines = []
    
    for line in lines:
        line_trimmed = line.strip()
        
        if line_trimmed.startswith("|"):
            in_table = True
            table_lines.append(line)
            continue
        else:
            if in_table:
                build_docx_table(doc, table_lines)
                table_lines = []
                in_table = False
                
        if line_trimmed.startswith('# '):
            add_docx_paragraph(doc, line_trimmed[2:], "H1")
        elif line_trimmed.startswith('## '):
            add_docx_paragraph(doc, line_trimmed[3:], "H2")
        elif line_trimmed.startswith('### '):
            add_docx_paragraph(doc, line_trimmed[4:], "H3")
        elif line_trimmed == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("____________________________________________________")
            run.font.color.rgb = RGBColor(200, 200, 200)
        elif not line_trimmed:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
        else:
            add_docx_paragraph(doc, line_trimmed, "Normal")
            
    if in_table and table_lines:
        build_docx_table(doc, table_lines)
        
    doc.save(output_docx_path)
    print(f"Word document saved successfully at: {output_docx_path}")

def find_today_reports():
    today_str = datetime.today().strftime('%Y-%m-%d')
    search_dirs = ['.', './불필요한 작업파일']
    domestic_md = None
    foreign_md = None
    
    for s_dir in search_dirs:
        if not os.path.exists(s_dir):
            continue
        for file in os.listdir(s_dir):
            if file.startswith(today_str) and file.endswith('.md'):
                if '취합' in file or '업로드' in file:
                    continue
                full_path = os.path.join(s_dir, file)
                if '해외' in file:
                    foreign_md = full_path
                else:
                    domestic_md = full_path
                    
    if not domestic_md or not foreign_md:
        print("Today's markdown files not fully found. Searching for most recent ones...")
        all_mds = []
        for s_dir in search_dirs:
            if not os.path.exists(s_dir):
                continue
            for file in os.listdir(s_dir):
                if file.endswith('.md'):
                    if '취합' in file or '업로드' in file:
                        continue
                    full_path = os.path.join(s_dir, file)
                    all_mds.append((full_path, os.path.getmtime(full_path)))
        
        all_mds.sort(key=lambda x: x[1], reverse=True)
        
        for path, _ in all_mds:
            if '해외' in path and not foreign_md:
                foreign_md = path
            elif '해외' not in path and not domestic_md:
                domestic_md = path
                
    return domestic_md, foreign_md

def markdown_to_html(md_content):
    lines = md_content.split('\n')
    html_lines = []
    in_table = False
    table_headers = []
    table_rows = []
    
    for line in lines:
        line_trimmed = line.strip()
        
        if line_trimmed.startswith('|'):
            in_table = True
            if re.match(r'^\s*\|?\s*:\s*[-=]+', line_trimmed) or '---' in line_trimmed:
                continue
            cols = [c.strip() for c in line_trimmed.split('|')[1:-1]]
            if not table_headers:
                table_headers = cols
            else:
                table_rows.append(cols)
            continue
        else:
            if in_table:
                table_html = "<table style='border-collapse: collapse; width: 100%; border: 1px solid #ddd; margin: 15px 0;'>"
                table_html += "<tr style='background-color: #f2f2f2; text-align: left;'><th style='border: 1px solid #ddd; padding: 10px; font-weight: bold;'>" + "</th><th style='border: 1px solid #ddd; padding: 10px; font-weight: bold;'>".join(table_headers) + "</th></tr>"
                for row in table_rows:
                    table_html += "<tr>"
                    for val in row:
                        val_cleaned = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2" target="_blank" style="color: #0073e6; text-decoration: none;">\1</a>', val)
                        table_html += f"<td style='border: 1px solid #ddd; padding: 10px;'>{val_cleaned}</td>"
                    table_html += "</tr>"
                table_html += "</table>"
                html_lines.append(table_html)
                in_table = False
                table_headers = []
                table_rows = []
                
        if line_trimmed.startswith('# '):
            html_lines.append(f"<h1 style='font-size: 24px; color: #111; border-bottom: 2px solid #222; padding-bottom: 8px; margin-top: 30px;'>{line_trimmed[2:]}</h1>")
        elif line_trimmed.startswith('## '):
            html_lines.append(f"<h2 style='font-size: 20px; color: #0066cc; margin-top: 25px; border-bottom: 1px solid #ddd; padding-bottom: 5px;'>{line_trimmed[3:]}</h2>")
        elif line_trimmed.startswith('### '):
            html_lines.append(f"<h3 style='font-size: 16px; color: #333; margin-top: 20px;'>{line_trimmed[4:]}</h3>")
        elif line_trimmed == '---':
            html_lines.append("<hr style='border: 0; height: 1px; background: #ccc; margin: 20px 0;'>")
        elif not line_trimmed:
            html_lines.append("<p>&nbsp;</p>")
        else:
            clean_line = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', line_trimmed)
            clean_line = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2" target="_blank" style="color: #0073e6; text-decoration: none;">\1</a>', clean_line)
            html_lines.append(f"<p style='line-height: 1.6; color: #444; margin: 8px 0;'>{clean_line}</p>")
            
    if in_table and table_headers:
        table_html = "<table style='border-collapse: collapse; width: 100%; border: 1px solid #ddd; margin: 15px 0;'>"
        table_html += "<tr style='background-color: #f2f2f2; text-align: left;'><th style='border: 1px solid #ddd; padding: 10px; font-weight: bold;'>" + "</th><th style='border: 1px solid #ddd; padding: 10px; font-weight: bold;'>".join(table_headers) + "</th></tr>"
        for row in table_rows:
            table_html += "<tr>"
            for val in row:
                val_cleaned = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2" target="_blank" style="color: #0073e6; text-decoration: none;">\1</a>', val)
                table_html += f"<td style='border: 1px solid #ddd; padding: 10px;'>{val_cleaned}</td>"
            table_html += "</tr>"
        table_html += "</table>"
        html_lines.append(table_html)
        
    return '\n'.join(html_lines)

def main():
    print("Searching for domestic and foreign box office report markdown files...")
    dom_path, for_path = find_today_reports()
    
    if not dom_path and not for_path:
        print("Error: No markdown files found to process.")
        return

    target_date = datetime.today().strftime('%Y-%m-%d')
    for path in [dom_path, for_path]:
        if path:
            basename = os.path.basename(path)
            match = re.match(r'^(\d{4}-\d{2}-\d{2})', basename)
            if match:
                target_date = match.group(1)
                break

    print(f"Target Date resolved for naming: {target_date}")
    
    dom_content = ""
    for_content = ""
    
    if dom_path:
        print(f"Reading domestic report: {dom_path}")
        with open(dom_path, 'r', encoding='utf-8') as f:
            dom_content = f.read()
            
    if for_path:
        print(f"Reading foreign report: {for_path}")
        with open(for_path, 'r', encoding='utf-8') as f:
            for_content = f.read()

    # 1. Generate unified markdown (취합_*.md)
    merged_md = ""
    if dom_content:
        merged_md += dom_content
    if for_content:
        if merged_md:
            merged_md += "\n\n---\n\n"
        merged_md += for_content

    merged_md_filename = f"취합_{target_date}_한국&해외 박스오피스.md"
    with open(merged_md_filename, 'w', encoding='utf-8') as f:
        f.write(merged_md)
    print(f"Temporary unified markdown created: {merged_md_filename}")

    # 2. Direct convert unified markdown to unified docx using python-docx (No MS Word required!)
    os.makedirs("Word_보고서", exist_ok=True)
    output_docx_path = os.path.join("Word_보고서", f"취합_{target_date}_한국&해외 박스오피스.docx")
    
    print("Converting unified markdown to Word document (.docx) via python-docx...")
    convert_md_to_docx(merged_md, output_docx_path)

    # 3. Generate structured Naver Blog HTML
    merged_html = "<div style='font-family: \"Nanum Gothic\", sans-serif; max-width: 800px; margin: 0 auto;'>"
    merged_html += f"<p style='color: #777; font-size: 14px;'>자동 생성 일시: {datetime.today().strftime('%Y-%m-%d %H:%M:%S')}</p>"
    
    if dom_content:
        merged_html += "<!-- Domestic Box Office -->"
        merged_html += markdown_to_html(dom_content)
        
    if for_content:
        merged_html += "<br><hr style='border: 0; border-top: 2px dashed #999; margin: 40px 0;'><br>"
        merged_html += "<!-- Foreign Box Office -->"
        merged_html += markdown_to_html(for_content)
        
    merged_html += "<br><hr style='border: 0; height: 1px; background: #ccc; margin: 30px 0;'>"
    merged_html += "<div style='background-color: #f9f9f9; padding: 15px; border-left: 4px solid #4caf50; font-size: 13px; color: #555;'>"
    merged_html += "<p style='margin: 0; font-weight: bold;'>⚠️ 사실관계 검증 및 출처 안내</p>"
    merged_html += "<p style='margin: 5px 0 0 0;'>본 보고서의 데이터는 영화진흥위원회 통합전산망(KOBIS) 및 해외 박스오피스 모조(Box Office Mojo) 등 공인된 데이터 소스를 기반으로 AI가 자동 수집 및 크로스체크하여 작성되었습니다. 교차 검증을 통해 데이터 누락 및 오류를 보완하였습니다.</p>"
    merged_html += "</div>"
    merged_html += "</div>"
    
    os.makedirs("HTML_원고", exist_ok=True)
    output_html_path = os.path.join("HTML_원고", f"업로드_{target_date}_한국&해외 박스오피스.html")
    with open(output_html_path, 'w', encoding='utf-8') as f:
        f.write(merged_html)
        
    print(f"Naver Blog HTML generated at: {output_html_path}")

if __name__ == "__main__":
    main()
